"""导出服务：聚合笔记引用 → 编号列表 / docx / BibTeX。

数据与格式分层（PRD 4.5）：只存结构化元数据（literature.json），
格式化是导出时的纯函数（citation_formats.py，用户实现）。
"""

from __future__ import annotations

import io

from .frontmatter import parse_frontmatter
from .literature import LiteratureEntry, load_literature
from .vault import default_vault_path

NOTE_DIR = "笔记"


def collect_references(
    note_id: str | None = None, collection_id: str | None = None
) -> list[LiteratureEntry]:
    """聚合笔记 frontmatter cites → 去重文献列表（按文献在 literature.json 中顺序）。

    note_id：只看该笔记（相对路径去 .md）；collection_id：只看该集合的笔记；缺省 = 全部。
    """
    vault = default_vault_path()
    note_dir = vault / NOTE_DIR
    if not note_dir.exists():
        return []

    lit_map = {entry.id: entry for entry in load_literature()}
    seen: set[str] = set()
    result: list[LiteratureEntry] = []

    for note_path in sorted(note_dir.glob("*.md")):
        rel = str(note_path.relative_to(vault)).replace("\\", "/")
        doc_id = rel[:-3]
        if note_id and doc_id != note_id:
            continue
        meta, _ = parse_frontmatter(note_path.read_text(encoding="utf-8"))
        if collection_id and meta.get("collection") != collection_id:
            continue
        for lit_id in meta.get("cites") or []:
            if lit_id in seen or lit_id not in lit_map:
                continue
            seen.add(lit_id)
            result.append(lit_map[lit_id])
    return result


def format_references(entries: list[LiteratureEntry], formatter) -> list[str]:
    """编号自动重排（[1] [2]...）+ 逐条格式化（formatter 为 citation_formats 纯函数）。"""
    return [f"[{i}] {formatter(entry)}" for i, entry in enumerate(entries, start=1)]


def build_references_docx(entries: list[LiteratureEntry], formatter, title: str = "参考文献") -> bytes:
    """python-docx 生成 .docx 字节流。"""
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    for i, entry in enumerate(entries, start=1):
        doc.add_paragraph(f"[{i}] {formatter(entry)}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_bibtex(entries: list[LiteratureEntry]) -> str:
    """BibTeX 序列化；key 自动生成（第一作者姓氏 + 年份），冲突加后缀。"""
    used: dict[str, int] = {}
    items: list[str] = []
    for entry in entries:
        key = _bibtex_key(entry)
        if key in used:
            used[key] += 1
            key = f"{key}{chr(ord('a') + used[key] - 1)}"
        else:
            used[key] = 0
        items.append(_format_bib_entry(entry, key))
    return "\n\n".join(items)


def _bibtex_key(entry: LiteratureEntry) -> str:
    first_family = entry.authors[0].get("family") if entry.authors else ""
    base = (first_family or "unknown").lower()
    return base + (str(entry.year) if entry.year else "")


def _format_bib_entry(entry: LiteratureEntry, key: str) -> str:
    authors = " and ".join(
        (f"{a['family']}, {a['given']}" if a.get("given") else a["family"])
        for a in entry.authors
    )
    fields = [
        f"title = {{{_bibtex_escape(entry.title)}}}",
        f"year = {entry.year}" if entry.year else None,
        f"author = {{{authors}}}" if authors else None,
        f"journal = {{{_bibtex_escape(entry.venue)}}}" if entry.venue else None,
        f"volume = {{{entry.volume}}}" if entry.volume else None,
        f"number = {{{entry.issue}}}" if entry.issue else None,
        f"pages = {{{entry.pages}}}" if entry.pages else None,
        f"doi = {{{entry.doi}}}" if entry.doi else None,
    ]
    body = ",\n  ".join(field for field in fields if field)
    return f"@article{{{key},\n  {body}\n}}"


def _bibtex_escape(text: str) -> str:
    """BibTeX 特殊字符转义（& % # _）。"""
    return (
        text.replace("&", "\\&").replace("%", "\\%").replace("#", "\\#").replace("_", "\\_")
    )
